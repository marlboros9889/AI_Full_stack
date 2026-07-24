from docx import Document
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path

out = Path(__file__).with_name("marinboy_script_utf8.docx")
doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Arial", Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after, color in [("Heading 1",20,20,6,"000000"),("Heading 2",16,18,6,"000000"),("Heading 3",14,16,4,"434343")]:
    s=doc.styles[name]; s.font.name="Arial"; s.font.size=Pt(size); s.font.bold=False
    s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
r=p.add_run("마린보이 살롱 포트폴리오 발표 대본"); r.font.name="Arial"; r.font.size=Pt(26)
p=doc.add_paragraph(); r=p.add_run("총 발표 시간 약 4분 45초 | 스마트폰 발표자용"); r.font.color.rgb=RGBColor.from_string("555555")

sections = [
("1. 소개","20초","안녕하세요. 마린보이는 고객의 예약 과정과 미용실 관리자의 운영 업무를 하나로 연결한 풀스택 예약 서비스입니다. 고객은 시술을 탐색하고 예약할 수 있으며, 관리자는 예약, 휴무일, 시술 메뉴를 한 화면에서 관리할 수 있습니다.","표지에서 고객과 관리자를 연결한 서비스라는 점을 먼저 강조합니다.","마린보이는 단순 예약 화면이 아니라 고객 경험과 운영 업무를 함께 개선한 서비스입니다."),
("2. 문제 정의","25초","기존 전화, 메신저, 수기 장부 방식은 고객이 가능한 시간을 바로 확인하기 어렵고, 관리자가 변경 사항을 여러 곳에서 확인해야 했습니다. 특히 동시에 같은 시간이 요청되거나 시술 시간이 겹치는 문제는 운영자의 확인만으로 완전히 방지하기 어려웠습니다.","가능 시간 확인의 어려움, 정보 분산, 중복 예약 위험을 차례로 짚습니다.","핵심 문제는 화면의 불편보다 예약 데이터의 충돌 가능성이었습니다."),
("3. 해결 방향","30초","고객이 먼저 시술 메뉴를 확인한 다음 날짜와 실제 예약 가능한 시간만 선택하도록 흐름을 설계했습니다. 예약 요청, 진행 중인 나의 예약, 이전 시술 이력도 각각 별도 화면으로 분리해 필요한 정보만 빠르게 확인하도록 구성했습니다.","메뉴 선택에서 예약까지의 흐름과 기능별 화면 분리를 설명합니다.","한 화면에 모든 정보를 넣기보다 목적별로 화면을 분리했습니다."),
("4. 기술 선택","30초","백엔드는 Spring Boot, 화면은 Thymeleaf와 Bootstrap 5, 데이터 접근은 MyBatis, 저장소는 Oracle XE를 사용했습니다. 무조건 유행하는 기술을 선택한 것이 아니라 예약 서비스의 빠른 초기 로딩과 서버 중심 검증에 적합한 SSR 구조를 선택했습니다.","기술 이름만 읽지 말고 왜 SSR을 선택했는지 설명합니다.","기술 선택의 기준은 유행이 아니라 서비스 특성과 제한된 리소스였습니다."),
("5. 고객 기능","45초","메인에서는 시술 메뉴를 큰 카드로 확인할 수 있고 월별 유효 예약 건수를 기준으로 헤어와 네일 인기 시술 TOP 5를 제공합니다. 메뉴를 선택하면 예약 화면으로 이동하며 오늘 남아 있는 빈 시간을 우선 확인할 수 있습니다. 로그인 후에는 진행 예약과 이전 시술 이력을 각각 관리할 수 있습니다.","고객 메인, TOP 5, 예약 화면을 순서대로 보여줍니다.","오늘 예약 가능한 시간을 먼저 보여줘 고객의 선택 과정을 줄였습니다."),
("6. 핵심 로직","45초","이 프로젝트의 핵심은 화면에서 불가능한 시간을 숨기는 것에 그치지 않습니다. 서버가 영업시간, 휴무일, 현재 시간, 시술 소요시간과 기존 예약을 실시간으로 계산합니다. 사용자가 JavaScript를 우회하거나 동시에 같은 시간을 요청해도 저장 직전에 다시 충돌을 검사해 중복 예약을 차단합니다.","천천히 말하며 UI와 서버의 이중 방어를 가장 강하게 강조합니다.","단순 View가 아니라 서버 사이드 시간 충돌 방어 로직을 탑재했습니다."),
("7. 관리자 기능","40초","관리자는 전체 예약을 단일 대시보드에서 조회하고 요청, 확정, 완료 상태를 처리합니다. 완료된 예약은 다시 취소할 수 없도록 상태 전이 제약을 적용했습니다. 휴무일과 시술 메뉴의 등록, 수정, 삭제, 이미지 관리와 2년 이전 종료 데이터 정리도 관리자 권한으로 제공합니다.","전체 예약, 휴무일, 메뉴 CRUD, 데이터 정리 순서로 보여줍니다.","수기 장부와 메신저 확인을 단일 관리자 화면으로 통합했습니다."),
("8. 결과","30초","고객 예약 CRUD, 관리자 운영 기능, Oracle DB 연동과 실제 브라우저 검증까지 완료했습니다. 고객에게는 가능한 시간만 보여줘 예약 과정을 단순화했고, 관리자에게는 예약과 운영 정보를 한 화면에서 처리할 수 있는 기준을 제공했습니다.","실제 실행 화면과 최종 성과 슬라이드를 보여줍니다.","빠른 초기 로딩, 충돌 방어, 운영 리소스 절감이 최종 성과입니다."),
("9. 마무리","20초","마린보이는 화면 구현에 그치지 않고 예약 데이터의 무결성과 운영 효율을 함께 고려한 프로젝트입니다. 향후에는 CSRF 보호, 자동화 통합 테스트, 알림 기능과 운영 통계를 보강해 실제 서비스 수준으로 발전시키겠습니다. 감사합니다.","마지막 문장은 화면을 보지 말고 청중을 보며 말합니다.","고객의 선택은 단순하게, 운영 데이터는 엄격하게 설계했습니다.")]
for title,time,script,point,key in sections:
    doc.add_heading(f"{title}  |  {time}", level=1)
    for label,text,bold in [("발표 대본",script,False),("화면/발표 포인트",point,False),("강조 문장",key,True)]:
        p=doc.add_paragraph(); x=p.add_run(label+"\n"); x.bold=True; y=p.add_run(text); y.bold=bold
doc.add_heading("발표 직전 체크", level=1)
for item in ["핵심 로직 슬라이드에서는 말의 속도를 늦춥니다.","기술 이름보다 선택 이유와 서비스 효과를 설명합니다.","화면을 읽지 말고 화면이 증명하는 기능을 설명합니다.","시간이 부족하면 기술과 결과 부분을 각각 5초 줄입니다."]:
    doc.add_paragraph(item, style="List Bullet")
doc.save(out)
print(out)
